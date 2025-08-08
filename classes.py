import os
from dotenv import load_dotenv
import google.generativeai as genai
from google.cloud.vision_v1 import types
from google.cloud import vision
from pdf2image import convert_from_path
import io
import json
from pathlib import Path
from PIL import Image
import subprocess
from docx import Document


def inicializar_credenciais():
    # Carregar variáveis do arquivo .env
    load_dotenv()
    
    # Configura a autenticação
    os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = os.getenv('GOOGLE_APPLICATION_CREDENTIALS')  

inicializar_credenciais()


class Documento:
    def __init__(self, arquivo_pdf=None):
        self.diretorio = arquivo_pdf
        self.tipo_documento = None
        self.num_documento = ''
        self.razao_social =None
        self.competencia = None
        self.data_arrecadacao = None
        self.tipo_documento = None
        self.cliente = None
        self.num_instalacao = None
        self.valor = None
        self.i_estadual = None
        self.cpf_cnpj = None
        self.tipo_lancamento = None
        self.plano_contas = None
        self.nome_da_fazenda = None
        self.nome_prestador_servicos = None
        self.cpf_cnpj_prestador = None

    def definir_atributos(self, dados_extraidos:dict):
        self.tipo_documento = dados_extraidos.get('tipo_documento', None)
        self.razao_social = dados_extraidos.get('razao_social', None)
        self.cpf_cnpj = dados_extraidos.get('cpf_cnpj', None)
        self.num_documento = dados_extraidos.get('numero_documento', '')
        if self.num_documento == None:
            self.num_documento = ''
        self.data_arrecadacao = dados_extraidos.get('data_arrecadacao', None)
        self.num_instalacao = dados_extraidos.get('num_instalacao', None)
        self.n_nota_fiscal = dados_extraidos.get('n_nota_fiscal', None)
        self.competencia = dados_extraidos.get('competencia', '')
        self.valor = dados_extraidos.get('valor_total', None)
        self.i_estadual = dados_extraidos.get('i_estadual', None)
        self.tipo_lancamento = dados_extraidos.get('tipo_lancamento', None)
        self.plano_contas = dados_extraidos.get('plano_contas', None)
        self.nome_da_fazenda = dados_extraidos.get('nome_da_fazenda', None)
        self.nome_prestador_servicos = dados_extraidos.get('nome_prestador_servicos', None)
        self.cpf_cnpj_prestador = dados_extraidos.get('cpf_cnpj_prestador', None)
        
    def _extrair_texto_doc(self):
        """
        Extrai texto de arquivos .doc e .docx.
        Para arquivos .doc, utiliza o LibreOffice para conversão temporária.
        """
        if self.arquivo.suffix == '.docx':
            documento = Document(self.arquivo)
            texto = "\n".join([paragrafo.text for paragrafo in documento.paragraphs])
        elif self.arquivo.suffix == '.doc':
            # Convertendo .doc para .docx usando o LibreOffice
            temp_path = self.arquivo.with_suffix('.docx')
            subprocess.run(['soffice', '--headless', '--convert-to', 'docx', str(self.arquivo), '--outdir', str(self.arquivo.parent)])
            documento = Document(temp_path)
            texto = "\n".join([paragrafo.text for paragrafo in documento.paragraphs])
            temp_path.unlink()  # Remove o arquivo convertido após o uso
        else:
            raise ValueError("Formato não reconhecido para extração.")
        return texto

    def extrair_texto_de_imagem(self) -> str:
        POPPLER_PATH = r"C:\Users\User\Downloads\poppler-22.04.0\Library\bin"
        try:
            arquivo_pdf = self.diretorio
    
            client = vision.ImageAnnotatorClient()

            if Path(arquivo_pdf).suffix == '.pdf':
                # Converte o PDF para uma imagem temporária (primeira página)
                img_convertida = convert_from_path(arquivo_pdf, 300, poppler_path=POPPLER_PATH)[0]
            elif Path(arquivo_pdf).suffix == ('.docx', '.doc'):
                img_convertida = self._extrair_texto_doc()
            else:
                img_convertida = Image.open(arquivo_pdf)

            # Converte a imagem para bytes em memória (sem salvar no disco)
            imgByteArr = io.BytesIO()
            img_convertida.save(imgByteArr, format='PNG')
            imgByteArr = imgByteArr.getvalue()

            # Envia os bytes da imagem para a API
            image = types.Image(content=imgByteArr)

            # Faz a solicitação de detecção de texto
            response = client.text_detection(image=image)

            if response.error.message:
                 raise Exception(f"Erro na API Google Vision: {response.error.message}")
            
            textos = response.text_annotations

            # Retorna o texto completo detectado
            if textos:
                return textos[0].description # str
            else:
                return None
        except Exception as e:
            print(f"Erro ao extrair texto da imagem: {e}")
            return None

    def resposta_ia(self, texto, prompt) -> str:
        api_key = os.getenv("API_KEY")
        if not api_key:
            raise ValueError("API_KEY não encontrada. Verifique o arquivo .env ou variáveis de ambiente.")
        
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-1.5-flash')

        arquivo_str = json.dumps(texto, indent=2)
        prompt = f'{prompt}\n\ndados: {arquivo_str}'
        response = model.generate_content(prompt)

        return response

    def tratar_dados(self, dados)-> dict:
        candidates = dados._result.candidates
        # Extrair o conteúdo JSON da resposta
        output_json_str = candidates[0].content.parts[0].text.strip().strip('```json').strip('```')
        try:
            dados = json.loads(output_json_str)
        except json.JSONDecodeError as e:
            print("Erro ao decodificar o JSON:", e)
            return None  # Adiciona um fallback em caso de erro

        return dados  # Retorna um dicionário

    @staticmethod
    def lancamento_agronota(dicionario):
        if isinstance(dicionario, dict):
            try:
                lancamento_formatado = {
                    'N° do documento*': dicionario['numero_documento'],
                    'Tipo do documento*': dicionario['tipo_documento'],
                    'CPF/CNPJ Cliente/Fornecedor*': dicionario['cpf_cnpj'],
                    'Valor do lançamento*': dicionario['valor'],
                    'Vencimento*': dicionario['vencimento'],
                    'Tipo de lançamento*': dicionario['tipo_lancamento'],
                    'IE ou Código da propriedade*': dicionario['i_estadual'],
                    'NIRF da propriedade': '',
                    'Plano de Contas*': dicionario['plano_contas'],
                    'Histórico*': dicionario['historico'],
                    'N° da Conta Bancária*': dicionario['num_conta'],
                }
            except Exception as e:
                print(f'Erro de lançamento agronota: {e}')
        return lancamento_formatado
    
    def renomear_documento(self, arquivo:Path, tipo_documento:str=None):
        if not arquivo.name or arquivo.name == 'None':
            return None
        else:
            try:
                obj_path = Path(arquivo)
                extencao = obj_path.suffix
                if self.competencia == None:
                    self.competencia = 'Nome não encontrado'
                if tipo_documento == 'energia_eletrica':
                    novo_nome = obj_path.with_name(f'{self.competencia}{extencao}')
                else:
                    novo_nome = obj_path.with_name(f'{self.tipo_documento}-{self.competencia}{extencao}')
                arquivo.rename(novo_nome)
                return novo_nome
            except Exception as e:
                return print(f'Erro ao renomear arquivo: {e}')